import streamlit as st
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from io import BytesIO


# ======================
# SESSION STATE
# ======================

if "paa_questions" not in st.session_state:
    st.session_state.paa_questions = []

if "competitors_raw" not in st.session_state:
    st.session_state.competitors_raw = []

if "competitors_enriched" not in st.session_state:
    st.session_state.competitors_enriched = []

if "title_tag" not in st.session_state:
    st.session_state.title_tag = ""

if "meta_description" not in st.session_state:
    st.session_state.meta_description = ""

if "article" not in st.session_state:
    st.session_state.article = ""

if "start_generation" not in st.session_state:
    st.session_state.start_generation = False

if "serp_ready" not in st.session_state:
    st.session_state.serp_ready = False

if "last_keyword" not in st.session_state:
    st.session_state.last_keyword = ""

if "last_country" not in st.session_state:
    st.session_state.last_country = "it"

if "last_language" not in st.session_state:
    st.session_state.last_language = "it"

if "last_num_results" not in st.session_state:
    st.session_state.last_num_results = 5


# ======================
# FUNZIONI
# ======================

def get_competitors(keyword: str, num_results: int, serper_key: str, hl: str, gl: str):
    url = "https://google.serper.dev/search"

    headers = {
        "X-API-KEY": serper_key,
        "Content-Type": "application/json"
    }

    competitors = []
    start = 0

    blocked_domains = [
        "youtube.com",
        "youtu.be",
        "tiktok.com",
        "pinterest.com",
        "facebook.com",
        "instagram.com"
    ]

    while len(competitors) < num_results and start < 40:
        payload = {
            "q": keyword,
            "gl": gl,
            "hl": hl,
            "num": 10,
            "start": start
        }

        response = requests.post(url, json=payload, headers=headers, timeout=30)
        response.raise_for_status()
        data = response.json()

        organic = data.get("organic", [])

        for item in organic:
            link = item.get("link")

            if not link:
                continue

            if any(domain in link for domain in blocked_domains):
                continue

            competitors.append({
                "title": item.get("title", ""),
                "link": link
            })

            if len(competitors) >= num_results:
                break

        start += 10

    return competitors[:num_results]


# ======================
# PAA DA SERPAPI
# ======================

def get_people_also_ask(keyword: str, serpapi_key: str, hl: str, gl: str):

    url = "https://serpapi.com/search.json"

    params = {
        "engine": "google",
        "q": keyword,
        "hl": hl,
        "gl": gl,
        "api_key": serpapi_key
    }

    questions = []

    try:
        response = requests.get(url, params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        paa_items = data.get("related_questions", [])

        for item in paa_items:
            question = item.get("question")
            if question:
                questions.append(question)

    except Exception:
        pass

    unique_questions = []
    seen = set()

    for q in questions:
        if q not in seen:
            seen.add(q)
            unique_questions.append(q)

    return unique_questions[:10]


# ======================
# SCRAPING PAGINA
# ======================

def fetch_page(url: str):
    try:
        resp = requests.get(
            url,
            timeout=15,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                )
            }
        )

        html = resp.text
        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        text = " ".join(soup.get_text().split())

        return html, text[:18000]

    except Exception:
        return "", ""


def extract_metadata(html: str):

    soup = BeautifulSoup(html, "html.parser")

    title = soup.title.string.strip() if soup.title and soup.title.string else ""

    h1_tag = soup.find("h1")
    h1 = h1_tag.get_text(strip=True) if h1_tag else ""

    meta_desc = ""
    meta = soup.find("meta", attrs={"name": "description"})

    if meta and "content" in meta.attrs:
        meta_desc = meta["content"].strip()

    return title, h1, meta_desc


# ======================
# GENERAZIONE ARTICOLO
# ======================

def generate_article(keyword: str, competitors: list, paa: list, openai_key: str, language: str):

    client = OpenAI(api_key=openai_key)

    merged = ""

    for i, comp in enumerate(competitors, start=1):

        merged += f"""
COMPETITOR {i}

URL: {comp['link']}

TITLE: {comp['html_title']}
H1: {comp['h1']}
META: {comp['meta_desc']}

CONTENUTO:
{comp['text']}

------------------------------------
"""

    paa_block = "\n".join([f"- {q}" for q in paa]) if paa else "Nessuna PAA disponibile."

    prompt = f"""
Sei un content writer SEO esperto. Siamo nel 2026.

Scrivi un contenuto SEO completo per la keyword:

{keyword}

Language code della ricerca: {language}

Il risultato deve contenere:

TITLE TAG (max 60 caratteri)

META DESCRIPTION (max 155 caratteri)

ARTICOLO HTML (800-1500 parole)

L'articolo deve essere scritto in HTML pronto per CMS.

Regole HTML:

- usa <h2> e <h3>
- usa <p>
- usa <ul> <ol>
- usa <strong>
- usa <table> se utile
- NON includere <html> <body>

Le PAA NON devono comparire come Q&A.

PAA INSIGHTS:
{paa_block}

COMPETITOR DATA:
{merged}

Formato output:

TITLE TAG:
...

META DESCRIPTION:
...

ARTICLE HTML:
...
"""

    response = client.chat.completions.create(
        model="gpt-5.4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )

    content = response.choices[0].message.content or ""

    title = ""
    meta = ""
    article = content

    if "TITLE TAG:" in content:
        parts = content.split("TITLE TAG:", 1)[1]

        if "META DESCRIPTION:" in parts:
            title = parts.split("META DESCRIPTION:", 1)[0].strip()

        if "ARTICLE HTML:" in parts:
            meta = parts.split("META DESCRIPTION:", 1)[1].split("ARTICLE HTML:", 1)[0].strip()
            article = parts.split("ARTICLE HTML:", 1)[1].strip()

    return title, meta, article


# ======================
# WORD EXPORT
# ======================

def create_word_file(title_tag, meta_description, article):

    doc = Document()

    doc.add_heading("Title Tag", level=2)
    doc.add_paragraph(title_tag)

    doc.add_heading("Meta Description", level=2)
    doc.add_paragraph(meta_description)

    doc.add_heading("HTML Article", level=2)
    doc.add_paragraph(article)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def reset_generation_outputs():

    st.session_state.competitors_enriched = []
    st.session_state.title_tag = ""
    st.session_state.meta_description = ""
    st.session_state.article = ""


# ======================
# SIDEBAR API
# ======================

st.sidebar.title("API Configuration")

SERPER_KEY = st.sidebar.text_input(
    "Serper.dev API Key",
    type="password"
)

SERPAPI_KEY = st.sidebar.text_input(
    "SerpAPI Key (People Also Ask)",
    type="password"
)

OPENAI_KEY = st.sidebar.text_input(
    "OpenAI API Key",
    type="password"
)


# ======================
# LOGO
# ======================

st.image(
    "https://YOUR_LOGO_URL/logo.png",
    width=220
)


# ======================
# UI
# ======================

st.title("SEO Article Generator")

keyword = st.text_input("Main keyword", value=st.session_state.last_keyword)

num_results = st.number_input(
    "Numero contenuti su cui fare scraping",
    min_value=1,
    max_value=20,
    value=st.session_state.last_num_results
)

country = st.text_input("Country code (gl)", value="it")

language = st.text_input("Language code (hl)", value="it")


col1, col2 = st.columns(2)

with col1:
    generate = st.button("Genera contenuto", use_container_width=True)

with col2:
    clear = st.button("Reset", use_container_width=True)


if generate:

    if not SERPER_KEY or not SERPAPI_KEY or not OPENAI_KEY:
        st.error("Inserisci tutte le API key nella sidebar.")
        st.stop()

    with st.spinner("Recupero SERP..."):

        competitors_raw = get_competitors(
            keyword=keyword,
            num_results=num_results,
            serper_key=SERPER_KEY,
            hl=language,
            gl=country
        )

        paa_questions = get_people_also_ask(
            keyword=keyword,
            serpapi_key=SERPAPI_KEY,
            hl=language,
            gl=country
        )

    st.session_state.competitors_raw = competitors_raw
    st.session_state.paa_questions = paa_questions

    st.session_state.start_generation = True
    st.rerun()


# ======================
# GENERAZIONE
# ======================

if st.session_state.start_generation:

    st.session_state.start_generation = False

    competitors = []

    for comp in st.session_state.competitors_raw:

        html, text = fetch_page(comp["link"])

        html_title, h1, meta_desc = extract_metadata(html)

        competitors.append({
            **comp,
            "html_title": html_title,
            "h1": h1,
            "meta_desc": meta_desc,
            "text": text
        })

    with st.spinner("Generazione articolo..."):

        title_tag, meta_description, article = generate_article(
            keyword=keyword,
            competitors=competitors,
            paa=st.session_state.paa_questions,
            openai_key=OPENAI_KEY,
            language=language
        )

    st.session_state.title_tag = title_tag
    st.session_state.meta_description = meta_description
    st.session_state.article = article

    st.rerun()


# ======================
# OUTPUT
# ======================

if st.session_state.article:

    st.subheader("SEO Metadata")

    st.write("**Title Tag**")
    st.write(st.session_state.title_tag)

    st.write("**Meta Description**")
    st.write(st.session_state.meta_description)

    st.subheader("Articolo HTML")

    st.code(st.session_state.article, language="html")

    word_file = create_word_file(
        st.session_state.title_tag,
        st.session_state.meta_description,
        st.session_state.article
    )

    st.download_button(
        label="Scarica Word",
        data=word_file,
        file_name="articolo_seo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
